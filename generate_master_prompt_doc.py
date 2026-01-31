from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_master_prompt_doc():
    doc = Document()

    # Style definitions
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    heading = doc.add_heading('NEXUS AI: Master Prompt Engineering Guide', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    p = doc.add_paragraph()
    p.add_run('This document contains a ').text
    run = p.add_run('Professional Master Prompt')
    run.bold = True
    p.add_run(' designed to instruct an Advanced AI Agent to build the NEXUS AI platform. It also includes a log of the key iterative prompts used during development.')

    doc.add_heading('🚀 Part 1: The Master Build Prompt', level=1)
    
    doc.add_paragraph('Copy and paste the following block into an advanced AI (Gemini 2.0, Claude 3.5 Sonnet) to recreate the project:')

    # Code Block (Simulated with Shading or Monospace font is hard in basic docx without styles, using Courier New)
    code_content = """### Role & Objective
You are a Principal Full-Stack AI Engineer specializing in Responsible AI and Fintech. Your task is to build NEXUS AI, a credit decisioning platform that prioritizes Fairness, Transparency, and Auditability.

### 🛠️ Tech Stack Constraints
*   Backend: Python (FastAPI) - must be async.
*   Frontend: Python (Streamlit) - must use a custom "Glassmorphism" CSS theme.
*   Database: SQLite (managed via SQLAlchemy).
*   ML Engine: Scikit-Learn (Random Forest) + Pandas.
*   Audit: Cryptographic hashing (SHA-256) for every decision.

### 🎨 Design System ("Pro Max" Style)
*   Theme: Dark Mode with Glassmorphism (translucent cards, blur effects).
*   Colors: Neon accents (Green for Approval, Red for Denial, Amber for Warning).
*   Typography: Clean, sans-serif (Inter/Roboto).
*   Charts: Use Donut Charts for "Approval Mix" and "Risk Profile".

### 📋 Core Functional Requirements

1.  The Credit Model:
    *   Train a Random Forest Classifier on the Indian Loan Dataset.
    *   CRITICAL: You MUST use a ColumnTransformer with OneHotEncoder for categorical variables.
    *   Output: Approval Status (Y/N) and Approval Probability (0-100%).

2.  The Audit Engine:
    *   Every prediction request MUST be logged to the SQLite database with a SHA-256 Hash.

3.  The Dashboard (Frontend):
    *   Summary Section: Big number metrics.
    *   Visualizations: Approval Mix (Donut), Risk Profile (Donut).
    *   Fairness Monitor: Calculate "Disparate Impact Ratio".
    *   Audit Trail: A searchable data table.

### 🧩 Step-by-Step Execution Plan

1.  Setup: Initialize project structure.
2.  ML Training: Write train.py (handle categoricals correctly!).
3.  Backend API: Create /predict and /audit endpoints.
4.  Frontend Layout: Apply st.markdown CSS for glassmorphism.
5.  Integration: Connect Frontend to Backend via HTTP.

Action: Generate the project structure and the core code files now."""

    p_code = doc.add_paragraph(code_content)
    p_code.style = 'No Spacing'
    for run in p_code.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 50, 0) # Dark Green for code

    doc.add_page_break()

    doc.add_heading('🔧 Part 2: Iterative Prompt Log (What Worked)', level=1)
    doc.add_paragraph('These are the specific prompts used to refine the product from MVP to "Pro Max".')

    # Iterative Prompts
    prompts = [
        ("1. Fixing the Model Logic", "Audit the current train.py file. I suspect it is ignoring categorical variables like 'Gender' and 'Education' because of a ValueError. Refactor the pipeline to use ColumnTransformer and OneHotEncoder so we use ALL available features."),
        ("2. UI/UX \"Glow Up\"", "The dashboard looks too basic. Upgrade the UI to a 'Pro Max' level using a Glassmorphism aesthetic. Inject Custom CSS for translucent cards and neon borders. Use high-contrast metric containers."),
        ("3. Visualizations Refinement", "Refine the Dashboard Visuals. Replace the generic volume area chart with two specific Donut Charts: 1. Approval Mix (Approved vs Denied) and 2. Risk Profile (Low/Medium/High risk buckets)."),
        ("4. Documentation Standard", "Generate a README.md that strictly follows these requirements: Problem Statement, Architecture Diagram, Tech Stack, AI Tools Used, Prompt Strategy, and Build Reproducibility Instructions."),
        ("5. Deployment & Cleanup", "Remove all temporary audit report files (*.txt) to clean up the repository. Then, force push the entire codebase to the remote origin main to ensure everything is perfectly synced.")
    ]

    for title, content in prompts:
        doc.add_heading(title, level=2)
        p_prompt = doc.add_paragraph(f"> \"{content}\"")
        p_prompt.style = 'Quote' # Tries to use default Quote style if available
        # Fallback manual styling for quote
        for run in p_prompt.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(80, 80, 80)

    doc.save('PROMPT_MASTER_CLASS.docx')
    print("Created PROMPT_MASTER_CLASS.docx successfully.")

if __name__ == "__main__":
    create_master_prompt_doc()
